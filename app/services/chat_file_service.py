"""
对话台文件上传与文本提取服务。

负责校验用户上传的简历或求职文档（PDF、DOCX、TXT、MD 等），
抽取清洗后的纯文本内容并生成元数据。
"""

from __future__ import annotations

import base64
import io
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional
import zipfile
from xml.etree import ElementTree

from fastapi import UploadFile
import requests

from app.core.config import settings
from app.core.text_sanitizer import sanitize_text

logger = logging.getLogger(__name__)

ALLOWED_CHAT_FILE_EXTENSIONS = frozenset(
    {
        ".pdf",
        ".docx",
        ".doc",
        ".txt",
        ".md",
        ".markdown",
        ".json",
        ".csv",
        ".yaml",
        ".yml",
        ".png",
        ".jpg",
        ".jpeg",
        ".webp",
        ".bmp",
    }
)
MAX_CHAT_FILE_SIZE = 20 * 1024 * 1024  # 20MB


class ChatFileService:
    """对话台文件与多模态图片解析服务。"""

    @classmethod
    async def process_upload(cls, upload_file: UploadFile) -> Dict[str, Any]:
        """校验并解析上传的文件或图片，返回纯文本与结构化元数据。"""
        original_name = upload_file.filename or "uploaded_file.txt"
        suffix = Path(original_name).suffix.lower()

        if suffix not in ALLOWED_CHAT_FILE_EXTENSIONS:
            raise ValueError(
                f"不支持的文件格式「{suffix}」，支持的格式包括：PDF、Word (.docx/.doc)、TXT、Markdown、JSON、CSV 以及图片 (PNG/JPG/WEBP/BMP)"
            )

        raw_bytes = await upload_file.read()
        file_size = len(raw_bytes)

        if file_size == 0:
            raise ValueError("上传的文件内容为空")
        if file_size > MAX_CHAT_FILE_SIZE:
            raise ValueError(f"文件大小超出限制（最大支持 20MB，当前为 {file_size / 1024 / 1024:.2f}MB）")

        # 提取文本
        mime_type = upload_file.content_type or ""
        extracted_text = cls.extract_text(raw_bytes, suffix, mime_type)
        cleaned_text = sanitize_text(extracted_text).strip()

        if not cleaned_text:
            raise ValueError("未能从文件中提取出有效文本，请确认文件未加密或损坏")

        char_count = len(cleaned_text)
        file_type_label = suffix.lstrip(".").upper()

        summary = f"成功提取 {file_type_label} 文件「{original_name}」，共 {char_count} 字"

        return {
            "filename": original_name,
            "file_type": file_type_label,
            "file_size": file_size,
            "char_count": char_count,
            "text": cleaned_text,
            "summary": summary,
        }

    @classmethod
    def extract_text(cls, raw_bytes: bytes, suffix: str, mime_type: str = "") -> str:
        """根据文件类型分支抽取文本。"""
        if suffix in (".png", ".jpg", ".jpeg", ".webp", ".bmp") or "image" in mime_type.lower():
            return cls._parse_image(raw_bytes, suffix, mime_type)
        if suffix == ".pdf" or "pdf" in mime_type.lower():
            return cls._parse_pdf(raw_bytes)
        if suffix in (".docx", ".doc") or "word" in mime_type.lower() or "officedocument" in mime_type.lower():
            return cls._parse_word(raw_bytes)
        return cls._decode_plain_text(raw_bytes)

    @classmethod
    def _parse_image(cls, raw_bytes: bytes, suffix: str, mime_type: str = "") -> str:
        """使用 SiliconFlow 多模态 OCR / 视觉大模型提取图片中的文本与视觉结构。"""
        fmt = suffix.lstrip(".").lower()
        if fmt == "jpg":
            fmt = "jpeg"
        if not fmt or fmt not in ("png", "jpeg", "webp", "bmp"):
            fmt = "png"

        b64_str = base64.b64encode(raw_bytes).decode("utf-8")
        data_uri = f"data:image/{fmt};base64,{b64_str}"

        api_key = settings.OPENAI_API_KEY or settings.SILICONFLOW_API_KEY
        if not api_key or "test" in api_key.lower() or "your-api-key" in api_key.lower():
            return f"【图片附件 {suffix.upper()}】（已接收图片数据，当前处于测试环境）"

        # 优先使用 OCR 专用模型，如 PaddleOCR 或 Qwen3-VL
        models_to_try = [
            settings.OCR_MODEL or "PaddlePaddle/PaddleOCR-VL-1.5",
            settings.VISION_MODEL or "Qwen/Qwen3-VL-8B-Instruct",
            "deepseek-ai/DeepSeek-OCR",
        ]

        url = f"{settings.OPENAI_API_BASE.rstrip('/')}/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        for model_name in models_to_try:
            try:
                payload = {
                    "model": model_name,
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": "请精确转录并提取图片中的所有文本内容，保持段落、标题与表格排版结构。如果图片是架构图或无文字图，请详细描述图中的核心组件、流程与设计细节。",
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {"url": data_uri},
                                },
                            ],
                        }
                    ],
                    "max_tokens": 2048,
                }
                resp = requests.post(url, headers=headers, json=payload, timeout=25)
                if resp.status_code == 200:
                    result_data = resp.json()
                    choices = result_data.get("choices", [])
                    if choices:
                        content = choices[0].get("message", {}).get("content", "").strip()
                        if content:
                            logger.info(f"多模态模型 {model_name} 成功解析图片，长度 {len(content)} 字符")
                            return content
            except Exception as e:
                logger.warning(f"多模态模型 {model_name} 解析图片异常: {e}")
                continue

        return f"【图片解析结果】成功读取图片 {suffix.upper()} 文件（未能提取出文字，已作为视觉上下文保留）。"

    @classmethod
    def _parse_pdf(cls, raw_bytes: bytes) -> str:
        """解析 PDF 文件文本。"""
        # 1. 尝试使用 pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw_bytes))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages_text.append(text.strip())
            if pages_text:
                return "\n\n".join(pages_text)
        except Exception as exc:
            logger.warning("pypdf parsing failed, trying PyPDF2: %s", exc)

        # 2. 尝试使用 PyPDF2
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(raw_bytes))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            if any(pages_text):
                return "\n\n".join(t.strip() for t in pages_text if t.strip())
        except Exception as exc:
            logger.warning("PyPDF2 parsing failed: %s", exc)

        # 3. 兜底纯文本
        return cls._decode_plain_text(raw_bytes)

    @classmethod
    def _parse_word(cls, raw_bytes: bytes) -> str:
        """解析 Word 文档文本。"""
        # 1. 尝试 python-docx
        try:
            from docx import Document

            doc = Document(io.BytesIO(raw_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # 同时读取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            if paragraphs:
                return "\n\n".join(paragraphs)
        except Exception as exc:
            logger.warning("python-docx parsing failed, trying OpenXML fallback: %s", exc)

        # 2. 尝试 OpenXML ZIP 解析
        try:
            return cls._parse_openxml_zip(raw_bytes)
        except Exception as exc:
            logger.warning("OpenXML parsing failed: %s", exc)

        # 3. 兜底纯文本
        return cls._decode_plain_text(raw_bytes)

    @classmethod
    def _parse_openxml_zip(cls, raw_bytes: bytes) -> str:
        """通过解析 OpenXML 内部 document.xml 抽取文本。"""
        if not zipfile.is_zipfile(io.BytesIO(raw_bytes)):
            return cls._decode_plain_text(raw_bytes)

        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
            if "word/document.xml" not in z.namelist():
                return cls._decode_plain_text(raw_bytes)

            xml_content = z.read("word/document.xml")
            tree = ElementTree.fromstring(xml_content)
            
            # 命名空间处理
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            text_nodes = tree.findall(".//w:t", ns)
            texts = [node.text for node in text_nodes if node.text]
            return "".join(texts)

    @classmethod
    def _decode_plain_text(cls, raw_bytes: bytes) -> str:
        """支持 UTF-8, GBK, GB18030, UTF-16 多编码自动兼容解码。"""
        if raw_bytes.startswith(b"\xef\xbb\xbf"):
            return raw_bytes.decode("utf-8-sig", errors="ignore")
        if raw_bytes.startswith(b"\xff\xfe") or raw_bytes.startswith(b"\xfe\xff"):
            return raw_bytes.decode("utf-16", errors="ignore")

        encodings = ("utf-8", "gb18030", "gbk", "big5", "cp936", "latin-1")
        for enc in encodings:
            try:
                return raw_bytes.decode(enc)
            except UnicodeDecodeError:
                continue

        return raw_bytes.decode("utf-8", errors="replace")
